# pdf_ingest.py
import os
import uuid
from typing import List
from dotenv import load_dotenv
load_dotenv()

import pypdf
from docx import Document
from openai import OpenAI
from supabase import create_client

EMBED_MODEL = os.getenv("EMBED_MODEL", "text-embedding-3-small")

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
supabase = create_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_SERVICE_ROLE_KEY")  # ローカルなら service role 推奨
)

def extract_text_from_pdf(pdf_path: str) -> str:
    reader = pypdf.PdfReader(pdf_path)
    texts = []
    for page in reader.pages:
        texts.append(page.extract_text() or "")
    return "\n".join(texts).strip()

def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()

def extract_text_from_txt(txt_path: str) -> str:
    for enc in ("utf-8", "cp932"):
        try:
            with open(txt_path, "r", encoding=enc) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()

def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    if ext == ".txt":
        return extract_text_from_txt(file_path)
    raise RuntimeError(f"未対応のファイル形式です: {ext}")

def chunk_text(text: str, size: int = 900, overlap: int = 150) -> List[str]:
    if not text:
        return []
    chunks = []
    i = 0
    while i < len(text):
        chunks.append(text[i:i+size])
        i += max(1, size - overlap)
    return chunks

def embed_texts(chunks: List[str]) -> List[List[float]]:
    res = client.embeddings.create(model=EMBED_MODEL, input=chunks)
    return [d.embedding for d in res.data]

def ingest_document(file_id: int, file_path: str, file_name: str):
    text = extract_text(file_path)
    if not text:
        raise RuntimeError("ファイルからテキストが抽出できませんでした（スキャンPDFや空ファイルの可能性）")

    chunks = chunk_text(text)
    vectors = embed_texts(chunks)

    # ✅ rag_chunks の列に合わせる：id / content / embedding のみ
    rows = []
    for c, v in zip(chunks, vectors):
        rows.append({
            "id": str(uuid.uuid4()),  # uuid
            "content": c,
            "embedding": v,
        })

    supabase.table("rag_chunks").insert(rows).execute()

    return {"ingested_chunks": len(chunks)}
